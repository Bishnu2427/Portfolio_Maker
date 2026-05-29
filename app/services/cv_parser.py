import os


class CVParser:
    @staticmethod
    def parse(file_path: str) -> str:
        if not file_path or not os.path.exists(file_path):
            return ''
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            return CVParser._parse_pdf(file_path)
        if ext in ('.docx', '.doc'):
            return CVParser._parse_docx(file_path)
        if ext == '.txt':
            return CVParser._parse_txt(file_path)
        return ''

    @staticmethod
    def _parse_pdf(path: str) -> str:
        # pdfplumber: layout-aware, preserves reading order and column structure
        try:
            import pdfplumber
            parts: list[str] = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text(x_tolerance=2, y_tolerance=3)
                    if text:
                        parts.append(text)
            result = '\n'.join(parts).strip()
            if result:
                return result
        except Exception as e:
            print(f'[CVParser] pdfplumber error: {e}')

        # PyPDF2 fallback
        try:
            import PyPDF2
            parts = []
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        parts.append(t)
            return '\n'.join(parts).strip()
        except Exception as e:
            print(f'[CVParser] PyPDF2 error: {e}')
            return ''

    @staticmethod
    def _parse_docx(path: str) -> str:
        try:
            import docx
            doc = docx.Document(path)
            lines: list[str] = []
            for para in doc.paragraphs:
                if para.text.strip():
                    lines.append(para.text)
            # Also extract text from tables (skills/education often in tables)
            for table in doc.tables:
                for row in table.rows:
                    row_text = '  '.join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        lines.append(row_text)
            return '\n'.join(lines).strip()
        except Exception as e:
            print(f'[CVParser] DOCX error: {e}')
            return ''

    @staticmethod
    def _parse_txt(path: str) -> str:
        try:
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read().strip()
        except Exception as e:
            print(f'[CVParser] TXT error: {e}')
            return ''
